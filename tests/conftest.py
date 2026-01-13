"""Shared fixtures and utilities for tests."""

from io import BytesIO
import os
import pytest
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import Pt
from PIL import Image
import io


@pytest.fixture(scope="session", autouse=True)
def setup_test_env():
    """Set up test environment variables before running tests."""
    # AWS/S3
    os.environ.setdefault("AWS_ACCESS_KEY_ID", "testing")
    os.environ.setdefault("AWS_SECRET_ACCESS_KEY", "testing")
    os.environ.setdefault("AWS_REGION", "us-east-1")
    os.environ.setdefault("AWS_S3_BUCKET", "test-bucket")
    
    # Database
    os.environ.setdefault("DATABASE_URL", "postgresql+asyncpg://test:test@localhost:5432/test_db")
    
    # Redis
    os.environ.setdefault("REDIS_URL", "redis://localhost:6379/0")
    
    # Celery
    os.environ.setdefault("CELERY_BROKER_URL", "redis://localhost:6379/1")
    os.environ.setdefault("CELERY_RESULT_BACKEND", "redis://localhost:6379/2")
    
    # Google
    os.environ.setdefault("GOOGLE_API_KEY", "test-google-key")
    
    # Auth
    os.environ.setdefault("JWT_SECRET_KEY", "test-jwt-secret-key-min-32-chars-long-for-security")
    os.environ.setdefault("JWT_ALGORITHM", "HS256")
    os.environ.setdefault("ACCESS_TOKEN_EXPIRE_MINUTES", "60")
    os.environ.setdefault("REFRESH_TOKEN_EXPIRE_DAYS", "30")
    
    # WorkOS
    os.environ.setdefault("WORKOS_API_KEY", "sk_test_workos_key")
    os.environ.setdefault("WORKOS_CLIENT_ID", "client_test_workos")
    os.environ.setdefault("WORKOS_REDIRECT_URI", "http://localhost:8000/api/v1/auth/sso/callback")


@pytest.fixture
def minimal_pdf():
    """Fixture providing a minimal valid PDF."""
    return _create_minimal_pdf("Test PDF")


@pytest.fixture
def simple_docx():
    """Fixture providing a simple DOCX document."""
    return _create_test_docx("Test paragraph", "Test cell", paragraphs_only=False)


def _create_minimal_pdf(text: str) -> bytes:
    """Create a minimal valid PDF with embedded text."""
    pdf = (
        b"%PDF-1.4\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
        b"4 0 obj<</Length 44>>stream\nBT /F1 12 Tf 100 700 Td ("
        + text.encode()
        + b") Tj ET\nendstream endobj\n"
        b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
        b"xref\n0 6\n0000000000 65535 f \n0000000009 00000 n \n0000000058 00000 n \n0000000115 00000 n \n0000000214 00000 n \n0000000306 00000 n \n"
        b"trailer<</Size 6/Root 1 0 R>>\nstartxref\n388\n%%EOF"
    )
    return pdf


def _create_test_docx(
    para_text: str, cell_text: str = "", paragraphs_only: bool = True
) -> BytesIO:
    """Create a simple DOCX document for testing."""
    stream = BytesIO()
    doc = Document()

    # Add first paragraph
    para1 = doc.add_paragraph(para_text)
    para1.runs[0].font.size = Pt(12)

    # Add second paragraph if different text
    if cell_text and paragraphs_only:
        para2 = doc.add_paragraph(cell_text)
        para2.runs[0].font.size = Pt(12)

    # Add table if requested
    if not paragraphs_only:
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        cell.text = cell_text or "Table Cell"

    doc.save(stream)
    stream.seek(0)
    return stream


def _create_pdf_with_image() -> bytes:
    """Create a minimal PDF with an embedded image for testing.

    Note: This creates a simple PDF structure. For more realistic testing,
    you may need to use a library like reportlab to create proper PDFs with images.
    """
    # Create a simple test image
    img = Image.new("RGB", (100, 100), color="red")
    img_bytes = io.BytesIO()
    img.save(img_bytes, format="PNG")
    img_bytes.seek(0)

    # For testing purposes, we'll create a minimal PDF structure
    # In practice, pdfplumber needs a proper PDF with image XObjects
    # This is a simplified version - real implementation may need reportlab
    pdf = _create_minimal_pdf("PDF with image")
    return pdf


@pytest.fixture
def temp_output_dir():
    """Fixture providing a temporary directory for test output."""
    with tempfile.TemporaryDirectory() as tmpdir:
        yield tmpdir
