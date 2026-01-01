import asyncio
import logging
from io import BytesIO
import os
from pathlib import Path
import re
from typing import Iterable, Union

import pdfplumber
from docx import Document
from PIL import Image


logger = logging.getLogger(__name__)


async def extract_text_from_pdf_stream(
    file_stream: Union[BytesIO, bytes, bytearray, memoryview],
) -> str:
    """Return normalized text from a PDF stream without blocking the event loop."""
    stream = _prepare_stream(file_stream)
    return await asyncio.to_thread(_extract_pdf_text_sync, stream)


async def extract_text_from_docx_stream(
    file_stream: Union[BytesIO, bytes, bytearray, memoryview],
) -> str:
    """Return normalized text from a DOCX stream, including table cells, without blocking the event loop."""
    stream = _prepare_stream(file_stream)
    return await asyncio.to_thread(_extract_docx_text_sync, stream)


def _extract_pdf_text_sync(file_stream: BytesIO) -> str:
    """Sync PDF parsing used behind the async wrapper."""
    chunks: list[str] = []

    with pdfplumber.open(file_stream) as pdf:
        for idx, page in enumerate(pdf.pages, start=1):
            try:
                page_text = page.extract_text() or ""
            except Exception as exc:  # pragma: no cover - pdfplumber internals
                logger.warning("Skipping unreadable PDF page %s: %s", idx, exc)
                continue
            if page_text:
                chunks.append(page_text)

    text = _normalize_text(chunks)
    if not text:
        logger.info("PDF contained no extractable text")
    return text


def _extract_docx_text_sync(file_stream: BytesIO) -> str:
    """Sync DOCX parsing used behind the async wrapper."""
    doc = Document(file_stream)
    text = _normalize_text(_iter_docx_text(doc))
    if not text:
        logger.info("DOCX contained no extractable text")
    return text


def _iter_docx_text(doc: Document) -> Iterable[str]:
    for para in doc.paragraphs:
        if para.text:
            yield para.text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    yield cell.text


def _normalize_text(chunks: Iterable[str]) -> str:
    """Trim, collapse whitespace, and join text chunks with stable line breaks."""
    cleaned: list[str] = []
    for chunk in chunks:
        stripped = chunk.strip()
        if stripped:
            cleaned.append(stripped)

    if not cleaned:
        return ""

    text = "\n".join(cleaned)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _prepare_stream(
    file_stream: Union[BytesIO, bytes, bytearray, memoryview],
) -> BytesIO:
    if isinstance(file_stream, (bytes, bytearray, memoryview)):
        stream = BytesIO(file_stream)
    else:
        stream = file_stream

    if stream.closed:
        raise ValueError("file_stream is closed")

    stream.seek(0)
    return stream


def extract_images_from_pdf(
    file_stream: BytesIO, output_dir: str = "extracted_images"
) -> list[str]:
    """Extract images from a PDF file and save them to disk.

    Args:
        file_stream: A BytesIO stream containing the PDF data.
        output_dir: Directory where extracted images will be saved. Defaults to "extracted_images".

    Returns:
        A list of file paths (strings) to the extracted and saved images.

    Example:
        >>> with open('document.pdf', 'rb') as f:
        ...     stream = BytesIO(f.read())
        ...     image_paths = extract_images_from_pdf(stream)
        ...     print(f"Saved {len(image_paths)} images")
    """
    image_paths: list[str] = []

    # Create output directory if it doesn't exist
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    with pdfplumber.open(file_stream) as pdf:
        image_counter = 0

        for page_idx, page in enumerate(pdf.pages, start=1):
            try:
                page_images = page.images or []
            except Exception as exc:  # pragma: no cover - pdfplumber internals
                logger.warning("Skipping unreadable PDF page %s: %s", page_idx, exc)
                continue

            for img_idx, img in enumerate(page_images):
                try:
                    # Extract image using pdfplumber's page.to_image() or direct bbox crop
                    bbox = (img["x0"], img["top"], img["x1"], img["bottom"])
                    cropped = page.crop(bbox)
                    pil_image = cropped.to_image(resolution=150).original

                    # Generate unique filename
                    image_counter += 1
                    filename = f"image_page{page_idx}_{img_idx + 1}.png"
                    filepath = output_path / filename

                    # Save image
                    pil_image.save(str(filepath), format="PNG")
                    image_paths.append(str(filepath))

                except Exception as exc:
                    logger.warning(
                        "Failed to extract image %d from page %d: %s",
                        img_idx + 1,
                        page_idx,
                        exc,
                    )
                    continue

    if not image_paths:
        logger.info("PDF contained no extractable images")
    else:
        logger.info("Extracted and saved %d images to %s", len(image_paths), output_dir)

    return image_paths
